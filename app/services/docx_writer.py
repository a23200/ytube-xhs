import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.services.platforms import PlatformAdapter

BODY_FONT = "PingFang SC"
FALLBACK_FONT = "Arial"
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(91, 103, 116)
ACCENT = RGBColor(30, 92, 122)


def _set_run_font(run: Any, *, name: str = BODY_FONT, size: float = 11, bold: bool = False, color: RGBColor = INK) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), FALLBACK_FONT)
    fonts.set(qn("w:hAnsi"), FALLBACK_FONT)
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:cs"), FALLBACK_FONT)
    fonts.set(qn("w:hint"), "eastAsia")
    language = rpr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        rpr.append(language)
    language.set(qn("w:val"), "zh-CN")
    language.set(qn("w:eastAsia"), "zh-CN")


def _set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "D8E0E5")
        borders.append(node)


def _set_table_geometry(table: Any, widths_dxa: tuple[int, ...]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_width)
    tbl_width.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_width.set(qn("w:type"), "dxa")
    table_grid = table._tbl.tblGrid
    for child in list(table_grid):
        table_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        table_grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = width
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FALLBACK_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FALLBACK_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FALLBACK_FONT)
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1


def _safe_title(post: Dict[str, Any], metadata: Dict[str, Any]) -> str:
    titles = post.get("titles", []) or []
    title = str(titles[0] if titles else metadata.get("title") or "平台文章").strip()
    return title or "平台文章"


def _body_paragraphs(post: Dict[str, Any]) -> list[str]:
    hook = str(post.get("hook") or "").strip()
    body = str(post.get("body") or "").strip()
    parts = [hook] if hook else []
    parts.extend(item.strip() for item in re.split(r"\n\s*\n|\n", body) if item.strip())
    return parts


def write_article_docx(
    metadata: Dict[str, Any],
    post: Dict[str, Any],
    platform: PlatformAdapter,
    output_path: Path,
) -> Path:
    doc = Document()
    _configure_document(doc)

    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    _set_run_font(header.add_run(f"{platform.name}发布稿"), size=9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True
    _set_run_font(title.add_run(_safe_title(post, metadata)), size=22, bold=True, color=ACCENT)

    generated_at = datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")
    table = doc.add_table(rows=3, cols=2)
    _set_table_geometry(table, (1944, 7416))
    rows = [
        ("目标平台", platform.name),
        ("生成时间", generated_at),
        ("内容来源", str(metadata.get("url") or "未提供")),
    ]
    for row, (label, value) in zip(table.rows, rows):
        for cell in row.cells:
            _set_cell_margins(cell)
            cell.vertical_alignment = 1
        label_p = row.cells[0].paragraphs[0]
        value_p = row.cells[1].paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        value_p.paragraph_format.space_after = Pt(0)
        _set_run_font(label_p.add_run(label), size=9.5, bold=True, color=MUTED)
        _set_run_font(value_p.add_run(value), size=9.5, color=INK)
    _set_table_borders(table)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)

    for index, paragraph_text in enumerate(_body_paragraphs(post)):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0 if index == 0 else 0.33)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.35
        _set_run_font(paragraph.add_run(paragraph_text), size=11.5, color=INK)

    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(12)
    source.paragraph_format.space_after = Pt(0)
    _set_run_font(
        source.add_run(
            f"来源：{metadata.get('author') or '未知作者'} | {metadata.get('title') or '未命名内容'} | {metadata.get('url') or '未提供 URL'}"
        ),
        size=9,
        color=MUTED,
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(footer.add_run("由 ytube-xhs 基于来源事实生成；发布前请人工复核事实、授权与平台规则。"), size=8, color=MUTED)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
