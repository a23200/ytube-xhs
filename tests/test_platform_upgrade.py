import json
import subprocess
import sys
import threading
import time
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.api import routes
from app.main import app
from app.schemas.models import ProjectCreate, ProjectStatus
from app.services import xhs_writer
from app.services.article_quality import build_data_concretization_report, evaluate_article_quality, find_subheadings
from app.services.docx_writer import write_article_docx
from app.services.platforms import get_platform, platform_keys, public_platform_capabilities
from app.services.runtime_store import ProjectPaths, ProjectStore, write_json
from app.services.task_manager import ChildProcessRegistry, TaskManager


def _valid_post(platform: str = "xhs") -> dict:
    return {
        "content_type": "原创文章",
        "target_audience": ["普通读者"],
        "titles": ["标题一", "标题二", "标题三", "标题四", "标题五"],
        "cover_text": "先别急着下结论",
        "hook": "看似只是省一步，结果最贵的代价，反而藏在这一步里。",
        "body": "很多人习惯先给答案，再去找证据。更稳妥的做法，是先核对来源和口径，再把复杂概念换成生活中的话。\n\n这样写出来的内容更容易看懂，也不需要靠照搬原话制造可信感。",
        "image_plan": [
            {
                "page": 1,
                "role": "cover",
                "caption": "先核对证据",
                "source_frame_time": None,
                "source_frame_path": None,
                "content_point": "先核对来源再表达观点",
            }
        ],
        "hashtags": ["#内容创作"],
        "publish_suggestion": "发布前复核事实与授权。",
        "source_disclaimer": "基于授权来源二次创作。",
        "platform": platform,
    }


def _content_assets(source_text: str = "来源材料强调先核对真实证据和统计口径，再用自己的话解释复杂概念。") -> dict:
    return {
        "one_sentence_summary": "先核对来源，再形成自己的判断。",
        "core_points": [
            {
                "point": "核对来源",
                "why_it_matters": "避免错误和照搬",
                "evidence": [{"type": "transcript", "time": 1.0, "text": source_text}],
            }
        ],
        "source_evidence": [
            {"claim": "核对来源", "source_type": "transcript", "time": 1.0, "source_text": source_text}
        ],
    }


def _registered_analysis(store: ProjectStore, project_id: str) -> None:
    paths = store.paths(project_id)
    files = {
        "metadata": (paths.source_dir / "metadata.json", {"title": "来源标题", "author": "作者", "url": "https://example.com/video"}),
        "transcript": (paths.transcript_dir / "transcript.json", {"segments": [{"start": 0, "end": 2, "text": "真实字幕", "source": "subtitle"}]}),
        "keyframes": (paths.analysis_dir / "keyframes.json", {"keyframes": [], "frame_count": 0}),
        "visual_analysis": (paths.analysis_dir / "visual-analysis.json", {"frames": []}),
        "content_assets": (
            paths.analysis_dir / "content-assets.json",
            {
                **_content_assets(),
                "golden_quotes": [{"quote": "改写金句", "rewrite_note": "已改写", "time": 1.0}],
                "chapters": [{"title": "段落", "summary": "总结", "start": 0, "end": 2}],
                "steps": [{"step": "核对", "evidence_time": 1.0}],
                "audience": ["读者"],
                "pain_points": ["信息难懂"],
                "xiaohongshu_angles": ["讲人话"],
                "recommended_content_type": "原创文章",
            },
        ),
    }
    for kind, (path, payload) in files.items():
        write_json(path, payload)
        store.add_output(project_id, kind, path)
    store.set_status(project_id, ProjectStatus.analysis_completed, "ready")


def test_platform_registry_exposes_four_truthful_adapters():
    assert platform_keys() == ("xhs", "toutiao", "douyin", "bilibili")
    capabilities = {item["key"]: item for item in public_platform_capabilities()}
    assert set(capabilities) == set(platform_keys())
    assert all(item["supports"]["docx_export"] for item in capabilities.values())
    assert all(item["supports"]["automatic_publish"] is False for item in capabilities.values())
    assert capabilities["douyin"]["supports"]["image_generation"] is False


@pytest.mark.parametrize("platform", platform_keys())
def test_platform_writer_uses_platform_prompt_and_artifact(tmp_path: Path, monkeypatch, platform: str):
    paths = ProjectPaths(tmp_path / platform)
    paths.ensure()
    seen = []

    def fake_chat(messages, **kwargs):
        seen.append(messages)
        return _valid_post(platform)

    monkeypatch.setattr(xhs_writer.llm_client, "json_chat", fake_chat)
    payload = xhs_writer.write_platform_post(
        {"title": "来源"},
        _content_assets(),
        {"keyframes": []},
        {"frames": []},
        "干货",
        paths,
        platform=platform,
    )

    adapter = get_platform(platform)
    assert payload["platform"] == platform
    assert adapter.name in seen[0][0]["content"]
    assert "不得出现任何小标题" in seen[0][0]["content"]
    assert (paths.analysis_dir / adapter.post_filename).exists()
    assert (paths.analysis_dir / adapter.quality_filename).exists()


def test_article_quality_detects_subheading_hook_and_verbatim_copy():
    source = "这是一段来源字幕中的连续原文片段，用来确认系统一定会识别长段复制而不是直接放行。"
    post = _valid_post()
    post["hook"] = "本文将介绍一个方法"
    post["body"] = f"一、背景\n{source}"
    report = evaluate_article_quality(post, _content_assets(source), platform="xhs")
    codes = {item["code"] for item in report["violations"]}
    assert {"mechanical_hook", "hook_lacks_contrast", "subheading_detected", "verbatim_source_copy_detected"} <= codes
    assert find_subheadings(post["body"])[0]["reason"] == "ordered_heading"


def test_data_concretization_requires_source_grounding():
    grounded = build_data_concretization_report(
        "调查显示 20% 的受访家庭遇到这个问题，总计约 12 万人。",
        "每 5 个受访家庭中，大约就有 1 个遇到问题，相当于约 12 万人。",
    )
    assert grounded["ungrounded_numeric_claims"] == []
    assert grounded["generated_population_expressions"][0]["source"]["people_equivalent"] == 120000

    converted_total = build_data_concretization_report(
        "调查覆盖 120,000 名用户，时间范围为 2025 年全年。",
        "调查覆盖约 12 万用户，时间范围仍为 2025 年全年。",
    )
    assert converted_total["ungrounded_numeric_claims"] == []
    assert converted_total["source_population_values"][0]["raw"] == "120,000 名用户"
    assert converted_total["generated_population_expressions"][0]["reason"] == "equivalent_source_population"

    fabricated = build_data_concretization_report("调查显示 20% 的受访家庭遇到这个问题。", "相当于约 12 万人。")
    assert fabricated["ungrounded_numeric_claims"][0]["reason"] == "source_total_or_exact_value_not_found"

    wrong_population = build_data_concretization_report("调查覆盖 12 万个家庭。", "相当于约 12 万用户。")
    assert wrong_population["ungrounded_numeric_claims"][0]["reason"] == "source_total_or_exact_value_not_found"


def test_docx_writer_creates_real_word_document_with_required_metadata(tmp_path: Path):
    output = tmp_path / "article.docx"
    write_article_docx(
        {"title": "来源标题", "author": "来源作者", "url": "https://example.com/video"},
        _valid_post("toutiao"),
        get_platform("toutiao"),
        output,
    )
    assert output.read_bytes().startswith(b"PK")
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "标题一" in text
    assert "看似只是省一步" in text
    assert "目标平台" in table_text and "今日头条" in table_text
    assert "内容来源" in table_text and "https://example.com/video" in table_text


def test_docx_file_route_returns_official_mime(tmp_path: Path, monkeypatch):
    store = ProjectStore(tmp_path)
    monkeypatch.setattr(routes, "store", store)
    record = store.create(ProjectCreate(url="https://example.com/video", text_only=True))
    paths = store.paths(record.project_id)
    output = paths.file_for_kind("douyin_post_docx")
    write_article_docx({"url": record.url}, _valid_post("douyin"), get_platform("douyin"), output)
    store.add_output(record.project_id, "douyin_post_docx", output)
    response = TestClient(app).get(f"/api/projects/{record.project_id}/files/douyin_post_docx")
    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert f'douyin-{record.project_id}-article.docx' in response.headers["content-disposition"]
    assert response.content.startswith(b"PK")


def test_partial_package_indexes_all_platform_article_artifacts(tmp_path: Path):
    from app.services.report_writer import write_partial_asset_package
    from app.services.runtime_store import ProjectPaths, write_json

    paths = ProjectPaths(tmp_path / "project")
    paths.ensure()
    for platform in ("douyin", "bilibili"):
        adapter = get_platform(platform)
        write_json(paths.analysis_dir / adapter.post_filename, {"platform": platform, "body": "正文"})
        write_json(paths.analysis_dir / adapter.quality_filename, {"passed": True})
        (paths.analysis_dir / adapter.markdown_filename).write_text("正文", encoding="utf-8")
        (paths.analysis_dir / adapter.docx_filename).write_bytes(b"PK-test")

    package = write_partial_asset_package(paths, {"code": "user_stopped"}, [])

    assert package["douyin_post"]["platform"] == "douyin"
    assert package["douyin_quality_report"]["passed"] is True
    assert package["bilibili_post"]["platform"] == "bilibili"
    assert package["bilibili_quality_report"]["passed"] is True
    assert package["available_files"]["douyin_post_markdown"].endswith("douyin-post.md")
    assert package["available_files"]["bilibili_post_docx"].endswith("bilibili-article.docx")


def test_platform_api_propagates_douyin_to_worker(tmp_path: Path, monkeypatch):
    store = ProjectStore(tmp_path)
    monkeypatch.setattr(routes, "store", store)
    monkeypatch.setattr(routes.llm_client, "ensure_available", lambda step: None)
    seen = []

    class ImmediateManager:
        def submit(self, scope, project_id, target, *args, platform=None, on_queued=None, **kwargs):
            seen.append((scope, project_id, args, platform))
            return {"queued": False, "queue_position": 0, "scope": scope, "platform": platform}

        def snapshot(self, project_id):
            return {"state": "idle", "queue_position": 0, "scope": None, "platform": None, "limits": {}}

        def cancel(self, project_id):
            return {"queued_cancelled": False, "running_cancelled": False, "terminated_child_processes": 0}

    monkeypatch.setattr(routes, "task_manager", ImmediateManager())
    record = store.create(ProjectCreate(url="https://example.com/video", text_only=True))
    _registered_analysis(store, record.project_id)
    response = TestClient(app).post(f"/api/projects/{record.project_id}/produce/platform/douyin")
    assert response.status_code == 200
    assert response.json()["platform"] == "douyin"
    assert seen == [("produce", record.project_id, ("douyin",), "douyin")]
    assert store.get(record.project_id).logs[-1].details["platform"] == "douyin"
    assert store.get(record.project_id).target_platform == "douyin"


def test_analyze_api_persists_selected_platform_end_to_end(tmp_path: Path, monkeypatch):
    store = ProjectStore(tmp_path)
    monkeypatch.setattr(routes, "store", store)
    seen = []

    class ImmediateManager:
        def submit(self, scope, project_id, target, *args, platform=None, on_queued=None, **kwargs):
            seen.append((scope, project_id, platform))
            return {"queued": False, "queue_position": 0, "scope": scope, "platform": platform}

        def snapshot(self, project_id):
            return {"state": "idle", "queue_position": 0, "scope": None, "platform": None, "limits": {}}

        def cancel(self, project_id):
            return {"queued_cancelled": False, "running_cancelled": False, "terminated_child_processes": 0}

    monkeypatch.setattr(routes, "task_manager", ImmediateManager())
    response = TestClient(app).post(
        "/api/projects/analyze",
        json={
            "url": "https://example.com/video",
            "target_platform": "bilibili",
            "text_only": True,
            "max_frames": 8,
        },
    )

    assert response.status_code == 200
    project_id = response.json()["project_id"]
    assert response.json()["target_platform"] == "bilibili"
    assert store.get(project_id).target_platform == "bilibili"
    assert seen == [("analyze", project_id, "bilibili")]
    status = TestClient(app).get(f"/api/projects/{project_id}/status").json()
    assert status["target_platform"] == "bilibili"
    assert status["progress"]["platform"] == "bilibili"


def test_legacy_project_record_infers_platform_then_persists_it(tmp_path: Path):
    store = ProjectStore(tmp_path)
    record = store.create(ProjectCreate(url="https://example.com/video"))
    status_path = store.paths(record.project_id).status_file()
    payload = json.loads(status_path.read_text(encoding="utf-8"))
    payload.pop("target_platform")
    payload["status"] = "toutiao_completed"
    status_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    migrated = store.get(record.project_id)
    assert migrated.target_platform == "toutiao"
    store.log(record.project_id, migrated.status, "legacy record loaded")
    persisted = json.loads(status_path.read_text(encoding="utf-8"))
    assert persisted["target_platform"] == "toutiao"


def test_task_manager_queues_and_cancel_releases_only_target_slot():
    manager = TaskManager(analyze_workers=1, produce_workers=2)
    release_a = threading.Event()
    release_b = threading.Event()
    started = {name: threading.Event() for name in ["a", "b", "c"]}
    completed = []

    def task(project_id: str, release: threading.Event):
        started[project_id].set()
        release.wait(timeout=2)
        completed.append(project_id)

    try:
        first = manager.submit("produce", "a", task, release_a)
        second = manager.submit("produce", "b", task, release_b)
        third = manager.submit("produce", "c", task, threading.Event())
        assert first["queued"] is False
        assert second["queued"] is False
        assert third["queued"] is True and third["queue_position"] == 1
        assert started["a"].wait(0.5) and started["b"].wait(0.5)

        cancelled = manager.cancel("a")
        assert cancelled["running_cancelled"] is True
        assert started["c"].wait(0.8)
        release_b.set()
        assert "b" in completed or started["b"].is_set()
        assert manager.snapshot("b")["state"] in {"running", "idle"}
    finally:
        release_a.set()
        release_b.set()
        manager.cancel("c")
        manager.shutdown()


def test_task_manager_cancel_queued_job_does_not_touch_running_job():
    manager = TaskManager(analyze_workers=1, produce_workers=1)
    release = threading.Event()
    running = threading.Event()
    second_started = threading.Event()

    def first(project_id: str):
        running.set()
        release.wait(timeout=2)

    def second(project_id: str):
        second_started.set()

    try:
        manager.submit("produce", "first", first)
        manager.submit("produce", "second", second)
        assert running.wait(0.5)
        cancelled = manager.cancel("second")
        assert cancelled["queued_cancelled"] is True
        assert manager.snapshot("first")["state"] == "running"
        release.set()
        time.sleep(0.05)
        assert not second_started.is_set()
    finally:
        release.set()
        manager.shutdown()


def test_child_process_registry_terminates_only_target_process_group():
    registry = ChildProcessRegistry()
    target = subprocess.Popen(
        [sys.executable, "-c", "import subprocess,time; subprocess.Popen(['sleep','30']); time.sleep(30)"],
        start_new_session=True,
    )
    other = subprocess.Popen([sys.executable, "-c", "import time; time.sleep(30)"], start_new_session=True)
    registry.register("target", target)
    registry.register("other", other)
    try:
        assert registry.terminate("target") == 1
        target.wait(timeout=2)
        assert target.returncode is not None
        assert other.poll() is None
    finally:
        registry.terminate("other")
        try:
            other.wait(timeout=2)
        except subprocess.TimeoutExpired:
            other.kill()


def test_child_process_registry_escalates_to_sigkill_for_ignored_sigterm():
    registry = ChildProcessRegistry()
    target = subprocess.Popen(
        [sys.executable, "-c", "import signal,time; signal.signal(signal.SIGTERM, signal.SIG_IGN); time.sleep(30)"],
        start_new_session=True,
    )
    registry.register("stubborn", target)
    try:
        time.sleep(0.1)
        assert registry.terminate("stubborn", grace_seconds=0.1) == 1
        target.wait(timeout=2)
        assert target.returncode is not None
    finally:
        if target.poll() is None:
            target.kill()
